from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document


def cell_text(cell) -> str:
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def iter_blocks(document: Document):
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            style = paragraph.style.name if paragraph.style is not None else ""
            yield f"[P|{style}] {text}"

    for table_index, table in enumerate(document.tables, start=1):
        yield f"\n[TABLE {table_index}]"
        for row in table.rows:
            values = [cell_text(cell) for cell in row.cells]
            yield "| " + " | ".join(values) + " |"


def main() -> None:
    parser = argparse.ArgumentParser(description="Extract paragraphs and tables from a docx file.")
    parser.add_argument("docx_path", type=Path)
    args = parser.parse_args()

    document = Document(args.docx_path)
    for block in iter_blocks(document):
        print(block)


if __name__ == "__main__":
    main()
